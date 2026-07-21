"""
Custom knowledge-base ingestion & retrieval.

Lets analysts upload arbitrary reference documents — playbooks, IR
runbooks, threat-intel reports, vendor advisories, past incident
write-ups, internal wikis exported to file, etc. — so the platform can
semantically search and (optionally) ground its reasoning against your
organization's own material. This mirrors the "drop files in, get a
searchable workspace knowledge base" workflow of tools like AnythingLLM,
built on the same on-prem ChromaDB instance THOS already uses for
hearth_kb/mitre_kb/siem_kb.

Storage: each document is split into overlapping word-chunks and
embedded into the 'custom_kb' Chroma collection (see
services/siem/clients.get_or_create_collection). Every chunk's metadata
carries enough to reconstruct/manage the parent document:
  doc_id        - stable id for the whole document (shared by all its chunks)
  filename      - original filename
  chunk_index   - 0-based position of this chunk within the document
  chunk_count   - total chunks for the parent document
  content_type  - file extension used to select the extractor
  ingested_at   - ISO-8601 UTC timestamp of ingestion

Supported formats: .txt, .md/.markdown, .csv/.tsv, .json, .log,
.html/.htm (tags stripped), .pdf (via pypdf), .docx (via python-docx).
Anything else is rejected with a clear error rather than being silently
ingested as garbage decoded-binary text.

Phase extension point: swap the naive word-based chunker for a
token-aware / sentence-boundary-aware splitter, or add OCR for scanned
PDFs, if recall on long technical documents needs improvement.
"""
import os
import io
import csv
import json
import uuid
import datetime
import logging

from services.siem.clients import get_or_create_collection

logger = logging.getLogger(__name__)

COLLECTION_NAME = "custom_kb"

# ~450 words (with 50-word overlap) approximates a several-hundred-token
# chunk without pulling in a tokenizer dependency just for this.
CHUNK_WORDS = 450
CHUNK_OVERLAP_WORDS = 50

# Generous cap for text/PDF/DOCX uploads; keeps one huge file from
# blocking the ingest tool call or ballooning the vector store.
MAX_DOCUMENT_BYTES = 25 * 1024 * 1024

SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".log",
    ".html", ".htm", ".pdf", ".docx",
}


class KnowledgeBaseError(Exception):
    """Raised for any user-facing ingestion problem (bad type, empty, too big, unparsable)."""


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_json(data: bytes) -> str:
    try:
        parsed = json.loads(data.decode("utf-8", errors="replace"))
        return json.dumps(parsed, indent=2)
    except json.JSONDecodeError:
        return data.decode("utf-8", errors="replace")


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    return "\n".join(", ".join(row) for row in reader)


def _extract_html(data: bytes) -> str:
    import re
    text = data.decode("utf-8", errors="replace")
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise KnowledgeBaseError(
            "PDF support requires the 'pypdf' package (add it to requirements.txt)"
        ) from e
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001 - one bad page shouldn't kill the whole doc
            logger.warning("failed to extract text from PDF page %d: %s", i, e)
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as e:
        raise KnowledgeBaseError(
            "DOCX support requires the 'python-docx' package (add it to requirements.txt)"
        ) from e
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


_EXTRACTORS = {
    ".txt": _extract_txt, ".log": _extract_txt,
    ".md": _extract_txt, ".markdown": _extract_txt,
    ".json": _extract_json,
    ".csv": _extract_csv, ".tsv": _extract_csv,
    ".html": _extract_html, ".htm": _extract_html,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


def _chunk_text(text: str, chunk_words: int = CHUNK_WORDS,
                 overlap: int = CHUNK_OVERLAP_WORDS) -> list[str]:
    words = text.split()
    if not words:
        return []
    chunks = []
    step = max(chunk_words - overlap, 1)
    for start in range(0, len(words), step):
        chunk = " ".join(words[start:start + chunk_words])
        if chunk.strip():
            chunks.append(chunk)
        if start + chunk_words >= len(words):
            break
    return chunks


def ingest_document(filename: str, content: bytes, doc_id: str | None = None) -> dict:
    """Extract, chunk, and embed a document into the custom_kb collection.

    Returns a summary dict on success. Raises KnowledgeBaseError (caught
    and turned into an {"error": ...} dict by the calling MCP tool) for
    any user-facing problem: unsupported type, empty/unreadable content,
    or oversized upload.
    """
    if not filename:
        raise KnowledgeBaseError("filename is required")
    if not content:
        raise KnowledgeBaseError(f"'{filename}' is empty")
    if len(content) > MAX_DOCUMENT_BYTES:
        raise KnowledgeBaseError(
            f"'{filename}' is too large ({len(content)} bytes) — max {MAX_DOCUMENT_BYTES} bytes"
        )

    ext = os.path.splitext(filename)[1].lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise KnowledgeBaseError(
            f"unsupported file type '{ext or '(none)'}' for '{filename}' — "
            f"supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    text = extractor(content).strip()
    if not text:
        raise KnowledgeBaseError(f"no extractable text found in '{filename}'")

    chunks = _chunk_text(text)
    if not chunks:
        raise KnowledgeBaseError(f"'{filename}' produced no chunks after processing")

    doc_id = doc_id or uuid.uuid4().hex[:12]
    ingested_at = datetime.datetime.utcnow().isoformat() + "Z"

    collection = get_or_create_collection(COLLECTION_NAME)
    ids = [f"{doc_id}::chunk::{i}" for i in range(len(chunks))]
    metas = [
        {
            "doc_id": doc_id, "filename": filename, "chunk_index": i,
            "chunk_count": len(chunks), "content_type": ext,
            "ingested_at": ingested_at,
        }
        for i in range(len(chunks))
    ]
    collection.upsert(ids=ids, documents=chunks, metadatas=metas)
    logger.info("ingested '%s' into custom_kb: doc_id=%s chunks=%d", filename, doc_id, len(chunks))

    return {
        "doc_id": doc_id, "filename": filename, "chunk_count": len(chunks),
        "char_count": len(text), "ingested_at": ingested_at,
    }


def list_documents() -> list[dict]:
    """Return one summary entry per ingested document (aggregated across its chunks)."""
    collection = get_or_create_collection(COLLECTION_NAME)
    if collection.count() == 0:
        return []
    res = collection.get(include=["metadatas"])
    by_doc: dict[str, dict] = {}
    for meta in res.get("metadatas", []):
        doc_id = meta.get("doc_id")
        if not doc_id or doc_id in by_doc:
            continue
        by_doc[doc_id] = {
            "doc_id": doc_id,
            "filename": meta.get("filename", ""),
            "chunk_count": meta.get("chunk_count", 0),
            "content_type": meta.get("content_type", ""),
            "ingested_at": meta.get("ingested_at", ""),
        }
    return sorted(by_doc.values(), key=lambda d: d.get("ingested_at", ""), reverse=True)


def delete_document(doc_id: str) -> dict:
    """Remove every chunk belonging to doc_id from the custom_kb collection."""
    collection = get_or_create_collection(COLLECTION_NAME)
    res = collection.get(where={"doc_id": doc_id}, include=[])
    ids = res.get("ids", [])
    if not ids:
        return {"deleted": False, "doc_id": doc_id, "chunks_removed": 0, "error": "document not found"}
    collection.delete(ids=ids)
    logger.info("deleted document from custom_kb: doc_id=%s chunks=%d", doc_id, len(ids))
    return {"deleted": True, "doc_id": doc_id, "chunks_removed": len(ids)}


def search(query: str, n_results: int = 5) -> list[dict]:
    """Semantic search over every ingested document's chunks."""
    collection = get_or_create_collection(COLLECTION_NAME)
    if collection.count() == 0 or not (query or "").strip():
        return []
    n_results = max(1, min(n_results, 20))
    res = collection.query(query_texts=[query], n_results=n_results)
    docs = res.get("documents", [[]])[0]
    metas = res.get("metadatas", [[]])[0]
    distances = (res.get("distances") or [[]])[0] if res.get("distances") else [None] * len(docs)
    return [{"text": d, "meta": m, "distance": dist} for d, m, dist in zip(docs, metas, distances)]
